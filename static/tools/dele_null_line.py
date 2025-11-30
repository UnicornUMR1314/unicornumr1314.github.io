from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def remove_empty_lines_from_docx(input_path, output_path):
    doc = Document(input_path)

    # 段落：删除空段落，非空段落合并连续换行
    paragraphs = list(doc.paragraphs)
    to_delete = []
    for idx, p in enumerate(paragraphs):
        text = p.text or ""
        if text.strip() == "":
            to_delete.append(idx)
        else:
            p.text = re.sub(r"\n+", "\n", text)

    for idx in reversed(to_delete):
        paragraphs[idx]._element.getparent().remove(paragraphs[idx]._element)

    # 表格：删除单元格内空段落，非空段落合并连续换行
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_paras = list(cell.paragraphs)
                del_idxs = []
                for i, cp in enumerate(cell_paras):
                    txt = cp.text or ""
                    if txt.strip() == "":
                        del_idxs.append(i)
                    else:
                        cp.text = re.sub(r"\n+", "\n", txt)
                for i in reversed(del_idxs):
                    cell_paras[i]._element.getparent().remove(cell_paras[i]._element)

    doc.save(output_path)
    print(f"空行清理完成！输出文件：{output_path}")

# ------------------- 调用示例 -------------------
import tkinter as tk
from tkinter import filedialog, messagebox


# 中文注释：通过图形界面选择输入与输出的 .docx 文件并执行清理
def main():
    try:
        root = tk.Tk()
        root.withdraw()

        input_path = filedialog.askopenfilename(
            title="选择输入 .docx 文件",
            filetypes=[("Word 文档", "*.docx")]
        )
        if not input_path:
            messagebox.showwarning("提示", "未选择输入文件，已取消。")
            return

        output_path = filedialog.asksaveasfilename(
            title="选择输出 .docx 文件",
            defaultextension=".docx",
            filetypes=[("Word 文档", "*.docx")]
        )
        if not output_path:
            messagebox.showwarning("提示", "未选择输出文件，已取消。")
            return

        remove_empty_lines_from_docx(input_path, output_path)
        messagebox.showinfo("完成", f"空行清理完成！\n输出文件：{output_path}")
    except Exception as e:
        try:
            messagebox.showerror("错误", f"处理失败：{e}")
        except Exception:
            print(f"处理失败：{e}")


if __name__ == "__main__":
    main()